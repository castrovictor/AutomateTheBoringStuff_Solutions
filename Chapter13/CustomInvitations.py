"""Write a program that would generate a Word document with custom
invitations that look like Figure 13-11."""

import docx, sys

guests = sys.argv[1]
invitations = sys.argv[2]

openFile = open(guests)
names = openFile.readlines()

doc = docx.Document(invitations)
for name in names:
    doc.add_paragraph('It would be a pleasure to have the company of',
                           style='Custom 1')
    doc.add_paragraph(name, style='Custom 2')
    doc.add_paragraph('at 11010 Memory Lane on the Evening of',
                           style='Custom 3')
    doc.add_paragraph('April 1st', style='Custom 4')
    doc.add_paragraph("at 7 o'clock", style='Custom 5')

    doc.add_page_break()

# Save document
document.save('invitationsGenerated.odt')
print("Invitations saved as invitationsGenerated.odt")
